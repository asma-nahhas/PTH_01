from django.shortcuts import redirect, render
from .models import Document
from .forms import DocumentForm
from django.http import JsonResponse
from django.http import HttpResponse
from docx import Document as DOCX
from django.core import serializers
from dict2xml import dict2xml
import json
import re
from Tables.models import Table



#Upload the Document file and save a copy to the database model Table

def uploadDocx(request):
    message =""
    # Handle file upload
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            newdoc = Document(docfile=request.FILES['docfile'])
            #newdoc.docfile.name='Table1.docx'
            newdoc.save()
            print(newdoc.docfile.name)

            #we should save all the Tables Paragraphs as table header
            Headers=[]
            f = open('media/'+newdoc.docfile.name, 'rb')
            doc=DOCX(f)

            fileIndexVar=0

            # Defining The Regular Expression of paraghraphs strings
            pattern = re.compile(r'^Table+[0-9]+$')


            #get All the pre added tables' headers
            allHeaders=[]
            allHeaders=Table.objects.values_list('name', flat=True)

            print("All Headers Are")
            print(allHeaders)

            print("All the Paraghraphs")

            allCurrentParags=[]
            for paragraph in doc.paragraphs:
                if paragraph.style.name=='Heading 1':
                    allCurrentParags.append(paragraph.text)

            print(allCurrentParags)

            for i,paragraph in enumerate(doc.paragraphs):
                if paragraph.style.name=='Heading 1':
                    
                    if re.match(pattern, paragraph.text):

                        print("Match found the table name"+paragraph.text)
                        message = 'File Matching the constraints Successfuly'

                        if paragraph.text in allHeaders or paragraph.text in allCurrentParags :
                            message='The Titles in this File is not Unique with previously added Titles'
                        else:
                            Headers.append(paragraph.text)
                            #Here We Should add this table title to the table Table
                            newTable=Table(name=paragraph.text,fileName=newdoc.docfile.name,fileIndex=fileIndexVar)
                            newTable.save();
                            fileIndexVar+=1;

                    else:
                        print("Match not found")
                        message = 'File headers mismatching the constraints'



            print (Headers)


            documents = Document.objects.all()

            # Redirect to the document list after POST
            return render(request,'list.html',{'documents': documents,'form': form,'message':message})
        else:
            message = 'The form is not valid. Fix the following error:'
    else:
        form = DocumentForm()  # An empty, unbound form

    # Load documents for the list page
    documents = Document.objects.all()

    # Render list page with the documents and the form
    context = {'documents': documents, 'form': form, 'message': message}
    return render(request, 'list.html', context)






  #Convert The selected table to a json object   

def  JsonTable(request):
    #f = open('media/documents/Table1.docx', 'rb')
    #document = DOCX(f)
    tblName =  request.POST["tblName"]
    print("The Table Name is : ")
    print(tblName)


    docFileName = Table.objects.filter(name=tblName)

    print(docFileName[0].fileName)

    print("Table Index is: "+ str(docFileName[0].fileIndex))

    #first we should get all the tables titles
    doc=DOCX('media/'+docFileName[0].fileName)

    #doc=DOCX('media/documents/demo.docx')

    tables = doc.tables
    
    table=doc.tables[docFileName[0].fileIndex]

    data = []
    title={}
    title["title"]=tblName


    data.append(title)

    #print(list(enumerate(table.rows)))

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)


        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            
            
            keys = tuple(text)
            
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        
        data.append(row_data)
        

    print(data)

    #Convert dictionary object to Json Object
    json_object = json.dumps(data, indent = 4)   

    #doc1 = Document.objects.filter(id=32)
    return render(request, 'demo.html', {"json":json_object})



#convert the selected table to a XML Object

def  XMLTable(request):

    tblName =  request.POST["tblName"]
    print("The Table Name is : ")
    print(tblName)

    docFileName = Table.objects.filter(name=tblName)

    print(docFileName[0].fileName)

    doc=DOCX('media/'+docFileName[0].fileName)
    tables = doc.tables
    print(tables)
    table=doc.tables[docFileName[0].fileIndex]
    data = []
  
    title={}
    title["title"]=tblName

    print(title)
    data.append(title)

    print("table is")
    print(table.rows[0].cells[0].text)
    print(list(enumerate(table.rows)))

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        print("the text is")
        print(text)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            print(i)
            
            keys = tuple(text)
            print(keys)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        print(text)
        data.append(row_data)
        print(data)

    print(data)

    #Convert the file to XML Object 
    xml = dict2xml(data) 
    print(xml)

    return render(request, 'xml.html', {"xml":xml})
